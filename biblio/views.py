from django import forms
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, FileResponse, HttpResponse, Http404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.core.paginator import Paginator
from django.db.models import Q, Count, Case, When, IntegerField, Value, F
from django.db.models.functions import Concat
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.utils.encoding import smart_str
import os
import mimetypes
from .models import Book, Author, Category, Publisher
from .forms import BookForm, AuthorForm, CategoryForm, PublisherForm
# Ajoutez ces imports en haut de views.py
from django.http import HttpResponse
from datetime import datetime
import io

# Pour Excel
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    openpyxl = None

# Pour PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    SimpleDocTemplate = None

# Pour Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None




# ============================================
# FONCTION UTILITAIRE POUR LES STATISTIQUES
# ============================================
def get_global_stats():
    """
    Récupère toutes les statistiques en une seule fois
    pour éviter les requêtes répétitives
    """
    # Statistiques des livres
    books_stats = Book.objects.aggregate(
        total=Count('book_id'),
        available=Count('book_id', filter=Q(status='available')),
        borrowed=Count('book_id', filter=Q(status='borrowed')),
        reserved=Count('book_id', filter=Q(status='reserved')),
        pdf_count=Count('book_id', filter=Q(file__iendswith='.pdf')),
        epub_count=Count('book_id', filter=Q(file__iendswith='.epub')),
        mobi_count=Count('book_id', filter=Q(file__iendswith='.mobi')),
    )
    
    # Comptage des autres entités
    total_authors = Author.objects.count()
    total_categories = Category.objects.count()
    total_publishers = Publisher.objects.count()
    
    # Langues disponibles
    languages = Book.objects.values('language').annotate(
        count=Count('book_id')
    ).order_by('-count')[:6]
    
    # Catégories avec comptage
    categories_with_counts = Category.objects.annotate(
        book_count=Count('book')
    ).order_by('-book_count')
    
    return {
        'stats': {
            'total': books_stats['total'],
            'available': books_stats['available'],
            'borrowed': books_stats['borrowed'],
            'reserved': books_stats['reserved'],
        },
        'total_authors': total_authors,
        'total_categories': total_categories,
        'total_publishers': total_publishers,
        'formats': {
            'pdf': books_stats['pdf_count'],
            'epub': books_stats['epub_count'],
            'mobi': books_stats['mobi_count'],
        },
        'languages': languages,
        'categories': categories_with_counts,
    }


# ============================================
# VUE INDEX
# ============================================
@login_required
def index(request):
    # Récupérer les paramètres de recherche et filtres
    search_query = request.GET.get('search', '').strip()
    category_id = request.GET.get('category', '')
    status_filter = request.GET.get('status', '')
    format_filter = request.GET.get('format', '')
    sort_by = request.GET.get('sort', 'created_at')
    
    # Base queryset pour les livres
    books_queryset = Book.objects.select_related('publisher').prefetch_related('authors', 'categories')
    
    # Appliquer la recherche
    if search_query:
        books_queryset = books_queryset.filter(
            Q(title__icontains=search_query) |
            Q(isbn__icontains=search_query) |
            Q(authors__name__icontains=search_query) |
            Q(publisher__publisher_name__icontains=search_query) |
            Q(description__icontains=search_query)
        ).distinct()
    
    # Appliquer le filtre de catégorie
    if category_id:
        books_queryset = books_queryset.filter(categories__category_id=category_id)
    
    # Appliquer le filtre de statut
    if status_filter:
        books_queryset = books_queryset.filter(status=status_filter)
    
    # Appliquer le filtre de format
    if format_filter == 'digital':
        books_queryset = books_queryset.exclude(file='').exclude(file__isnull=True)
    elif format_filter == 'physical':
        books_queryset = books_queryset.filter(Q(file='') | Q(file__isnull=True))
    
    # Appliquer le tri
    sort_mapping = {
        'title': 'title',
        'publication_year': '-publication_year',
        'created_at': '-created_at',
        '-title': '-title',
    }
    books_queryset = books_queryset.order_by(sort_mapping.get(sort_by, '-created_at'))
    
    # Livres populaires (affichés séparément en haut)
    popular_books = Book.objects.annotate(
        availability_score=Case(
            When(status='available', then=3),
            When(status='reserved', then=2),
            When(status='borrowed', then=1),
            default=0,
            output_field=IntegerField(),
        )
    ).order_by('-availability_score', '-created_at')[:10]
    
    # Statistiques pour les livres filtrés
    filtered_total = books_queryset.count()
    filtered_available = books_queryset.filter(status='available').count()
    
    # Récupérer les statistiques globales
    global_stats = get_global_stats()
    
    context = {
        'popular_books': popular_books,
        'all_books': books_queryset,
        'filtered_total': filtered_total,
        'filtered_available': filtered_available,
        # Paramètres de recherche actuels
        'current_search': search_query,
        'current_category': category_id,
        'current_status': status_filter,
        'current_format': format_filter,
        'current_sort': sort_by,
    }
    
    # Ajouter les statistiques globales
    context.update(global_stats)
    
    return render(request, 'biblio/index.html', context)


# ============================================
# API ENDPOINTS
# ============================================
@csrf_exempt
@require_http_methods(["GET"])
def api_books(request):
    page = request.GET.get('page', 1)
    per_page = request.GET.get('per_page', 12)
    search = request.GET.get('search', '')
    category_id = request.GET.get('category', None)
    status = request.GET.getlist('status', [])
    formats = request.GET.getlist('format', [])
    languages = request.GET.getlist('language', [])
    availability = request.GET.get('availability', None)
    sort_by = request.GET.get('sort', 'created_at')
    sort_order = request.GET.get('order', 'desc')
    min_year = request.GET.get('min_year', None)
    max_year = request.GET.get('max_year', None)
    
    books = Book.objects.all()
    
    if search:
        books = books.filter(
            Q(title__icontains=search) |
            Q(summary__icontains=search) |
            Q(authors__name__icontains=search) |
            Q(isbn__icontains=search) |
            Q(publisher__publisher_name__icontains=search)
        ).distinct()
    
    if category_id:
        books = books.filter(categories__category_id=category_id)
    
    if status:
        books = books.filter(status__in=status)
    
    if formats:
        format_filters = Q()
        for fmt in formats:
            format_filters |= Q(file__endswith=f'.{fmt}')
        books = books.filter(format_filters)
    
    if languages:
        books = books.filter(language__in=languages)
    
    if availability == 'available':
        books = books.filter(available_copies__gt=0)
    elif availability == 'unavailable':
        books = books.filter(available_copies=0)
    
    if min_year:
        books = books.filter(publication_year__gte=min_year)
    if max_year:
        books = books.filter(publication_year__lte=max_year)
    
    sort_field = sort_by
    if sort_order == 'desc':
        sort_field = f'-{sort_field}'
    
    if sort_by == 'title':
        books = books.order_by(sort_field)
    elif sort_by == 'author':
        books = books.annotate(
            author_name=Concat('authors__name')
        ).order_by(f'{sort_field}')
    elif sort_by == 'popularity':
        books = books.annotate(
            borrow_count=Count('borrowings')
        ).order_by(f'-borrow_count')
    elif sort_by == 'year':
        books = books.order_by(f'{sort_field}publication_year')
    else:
        books = books.order_by(sort_field)
    
    paginator = Paginator(books, per_page)
    try:
        page_obj = paginator.page(page)
    except:
        page_obj = paginator.page(1)
    
    total_available = books.filter(available_copies__gt=0).count()
    
    return JsonResponse({
        'books': [book.to_dict() for book in page_obj],
        'total': paginator.count,
        'available': total_available,
        'pages': paginator.num_pages,
        'current_page': page_obj.number,
        'has_next': page_obj.has_next(),
        'has_prev': page_obj.has_previous(),
        'filters_applied': {
            'search': bool(search),
            'category': bool(category_id),
            'status': bool(status),
            'formats': bool(formats),
            'languages': bool(languages),
            'year_range': bool(min_year or max_year),
        }
    })


@csrf_exempt
@require_http_methods(["POST"])
def api_create_book(request):
    form = BookForm(request.POST, request.FILES)
    if form.is_valid():
        book = form.save()
        return JsonResponse(book.to_dict(), status=201)
    else:
        return JsonResponse({'error': form.errors}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_get_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    return JsonResponse(book.to_dict())


@csrf_exempt
@require_http_methods(["PUT"])
def api_update_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    form = BookForm(request.POST, request.FILES, instance=book)
    
    if form.is_valid():
        book = form.save()
        return JsonResponse(book.to_dict())
    else:
        return JsonResponse({'error': form.errors}, status=400)


@csrf_exempt
@require_http_methods(["DELETE"])
def api_delete_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    
    try:
        if book.file:
            file_path = os.path.join(settings.MEDIA_ROOT, book.file.name)
            if os.path.exists(file_path):
                os.remove(file_path)
        
        book.delete()
        return JsonResponse({'message': 'Book deleted successfully'})
    
    except Exception as e:
        return JsonResponse({'error': f'Deletion error: {str(e)}'}, status=500)


def api_download_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    if not book.file:
        return HttpResponse('File not found', status=404)
    
    file_path = book.file.path
    if not os.path.exists(file_path):
        return HttpResponse('File not found', status=404)
    
    response = FileResponse(open(file_path, 'rb'))
    response['Content-Disposition'] = f'attachment; filename="{book.file.name}"'
    return response


@csrf_exempt
@require_http_methods(["GET"])
def api_authors(request):
    authors = Author.objects.all().order_by('name')
    return JsonResponse([author.to_dict() for author in authors], safe=False)


@csrf_exempt
@require_http_methods(["GET"])
def api_categories(request):
    categories = Category.objects.all().order_by('category_name')
    return JsonResponse([cat.to_dict() for cat in categories], safe=False)


@csrf_exempt
@require_http_methods(["GET"])
def api_publishers(request):
    publishers = Publisher.objects.all().order_by('publisher_name')
    return JsonResponse([pub.to_dict() for pub in publishers], safe=False)


# ============================================
# VUES AUTEURS
# ============================================
@login_required
@require_http_methods(["GET", "POST"])
def add_author(request):
    if request.method == "POST":
        form = AuthorForm(request.POST)
        if form.is_valid():
            author = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'author': author.to_dict()
                })
            messages.success(request, 'Auteur ajouté avec succès.')
            return redirect('add_author')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors.get_json_data()
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = AuthorForm()
    
    authors = Author.objects.all().order_by('name')
    
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'form_html': '...'})
    
    # Ajouter les statistiques globales
    context = {
        'form': form,
        'authors': authors
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/author_form.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def edit_author(request, author_id):
    author = get_object_or_404(Author, pk=author_id)
    if request.method == "POST":
        form = AuthorForm(request.POST, instance=author)
        if form.is_valid():
            author = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'author': author.to_dict()
                })
            messages.success(request, 'Auteur modifié avec succès.')
            return redirect('add_author')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = AuthorForm(instance=author)
    
    authors = Author.objects.all().order_by('name')
    
    context = {
        'form': form,
        'author': author,
        'authors': authors
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/author_form.html', context)


@login_required
@require_http_methods(["POST"])
def delete_author(request, author_id):
    author = get_object_or_404(Author, pk=author_id)
    try:
        author_name = f"{author.name}"
        author.delete()
        messages.success(request, f'Auteur "{author_name}" supprimé avec succès.')
    except Exception as e:
        messages.error(request, f'Erreur lors de la suppression: {str(e)}')
    
    return redirect('add_author')


# ============================================
# VUES CATÉGORIES
# ============================================
@login_required
@require_http_methods(["GET", "POST"])
def add_category(request):
    if request.method == "POST":
        form = CategoryForm(request.POST)
        if form.is_valid():
            category = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'category': category.to_dict()
                })
            messages.success(request, 'Catégorie ajoutée avec succès.')
            return redirect('add_category')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = CategoryForm()
    
    categories = Category.objects.all().order_by('category_name')
    
    context = {
        'form': form,
        'categories': categories
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/category_form.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def edit_category(request, category_id):
    category = get_object_or_404(Category, pk=category_id)
    if request.method == "POST":
        form = CategoryForm(request.POST, instance=category)
        if form.is_valid():
            category = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'category': category.to_dict()
                })
            messages.success(request, 'Catégorie modifiée avec succès.')
            return redirect('add_category')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = CategoryForm(instance=category)
    
    categories = Category.objects.all().order_by('category_name')
    
    context = {
        'form': form,
        'category': category,
        'categories': categories
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/category_form.html', context)


@login_required
@require_http_methods(["POST"])
def delete_category(request, category_id):
    category = get_object_or_404(Category, pk=category_id)
    try:
        category_name = category.category_name
        if category.book_set.exists():
            messages.error(request, f'Impossible de supprimer la catégorie "{category_name}" car elle est utilisée par des livres.')
        else:
            category.delete()
            messages.success(request, f'Catégorie "{category_name}" supprimée avec succès.')
    except Exception as e:
        messages.error(request, f'Erreur lors de la suppression: {str(e)}')
    
    return redirect('add_category')


# ============================================
# VUES ÉDITEURS
# ============================================
@login_required
@require_http_methods(["GET", "POST"])
def add_publisher(request):
    if request.method == "POST":
        form = PublisherForm(request.POST)
        if form.is_valid():
            publisher = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'publisher': publisher.to_dict()
                })
            messages.success(request, 'Éditeur ajouté avec succès.')
            return redirect('add_publisher')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = PublisherForm()
    
    publishers = Publisher.objects.all().order_by('publisher_name')
    
    context = {
        'form': form,
        'publishers': publishers
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/publisher_form.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def edit_publisher(request, publisher_id):
    publisher = get_object_or_404(Publisher, pk=publisher_id)
    if request.method == "POST":
        form = PublisherForm(request.POST, instance=publisher)
        if form.is_valid():
            publisher = form.save()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'success',
                    'publisher': publisher.to_dict()
                })
            messages.success(request, 'Éditeur modifié avec succès.')
            return redirect('add_publisher')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'errors': form.errors
                }, status=400)
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = PublisherForm(instance=publisher)
    
    publishers = Publisher.objects.all().order_by('publisher_name')
    
    context = {
        'form': form,
        'publisher': publisher,
        'publishers': publishers
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/publisher_form.html', context)


@login_required
@require_http_methods(["POST"])
def delete_publisher(request, publisher_id):
    publisher = get_object_or_404(Publisher, pk=publisher_id)
    try:
        publisher_name = publisher.publisher_name
        if publisher.book_set.exists():
            messages.error(request, f'Impossible de supprimer l\'éditeur "{publisher_name}" car il est utilisé par des livres.')
        else:
            publisher.delete()
            messages.success(request, f'Éditeur "{publisher_name}" supprimé avec succès.')
    except Exception as e:
        messages.error(request, f'Erreur lors de la suppression: {str(e)}')
    
    return redirect('add_publisher')


# ============================================
# VUES LIVRES
# ============================================
@login_required
@require_http_methods(["GET", "POST"])
def add_book(request):
    if request.method == "POST":
        form = BookForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                book = form.save(commit=False)
                book_type = request.POST.get('book_type', 'physical')
                
                if book_type == 'digital' and not book.file:
                    form.add_error('file', 'Un fichier est obligatoire pour les livres numériques')
                    raise forms.ValidationError("Fichier manquant pour livre numérique")
                
                elif book_type == 'physical':
                    book.file = None
                
                book.save()
                
                authors = request.POST.getlist('authors')
                categories = request.POST.getlist('categories')
                
                if authors:
                    book.authors.set(authors)
                if categories:
                    book.categories.set(categories)
                
                messages.success(request, 'Livre ajouté avec succès.')
                return redirect('book_list')
                
            except Exception as e:
                messages.error(request, f'Erreur lors de l\'ajout: {str(e)}')
        else:
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = BookForm()
    
    authors = Author.objects.all().order_by('name')
    categories = Category.objects.all().order_by('category_name')
    publishers = Publisher.objects.all().order_by('publisher_name')
    
    context = {
        'form': form,
        'authors': authors,
        'categories': categories,
        'publishers': publishers
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/book_form.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def edit_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    
    if request.method == "POST":
        form = BookForm(request.POST, request.FILES, instance=book)
        if form.is_valid():
            try:
                book = form.save(commit=False)
                book_type = request.POST.get('book_type', 'physical')
                
                if book_type == 'digital':
                    if not book.file and 'file' not in request.FILES:
                        form.add_error('file', 'Un fichier est obligatoire pour les livres numériques')
                        raise forms.ValidationError("Fichier manquant pour livre numérique")
                
                elif book_type == 'physical':
                    if book.file:
                        book.file.delete(save=False)
                        book.file = None
                
                book.save()
                
                authors = request.POST.getlist('authors')
                categories = request.POST.getlist('categories')
                
                if authors:
                    book.authors.set(authors)
                if categories:
                    book.categories.set(categories)
                
                messages.success(request, 'Livre modifié avec succès.')
                return redirect('book_list')
                
            except Exception as e:
                messages.error(request, f'Erreur lors de la modification: {str(e)}')
        else:
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = BookForm(instance=book)
    
    authors = Author.objects.all().order_by('name')
    categories = Category.objects.all().order_by('category_name')
    publishers = Publisher.objects.all().order_by('publisher_name')
    
    context = {
        'form': form,
        'book': book,
        'authors': authors,
        'categories': categories,
        'publishers': publishers
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/forms/book_form.html', context)


@login_required
@require_http_methods(["POST"])
def delete_book(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    try:
        book_title = book.title
        book.delete()
        messages.success(request, f'Livre "{book_title}" supprimé avec succès.')
    except Exception as e:
        messages.error(request, f'Erreur lors de la suppression: {str(e)}')
    
    return redirect('index')


@login_required
def book_list(request):
    books = Book.objects.all().select_related('publisher').prefetch_related('authors', 'categories')
    
    search = request.GET.get('search', '')
    category_id = request.GET.get('category', '')
    status = request.GET.get('status', '')
    
    if search:
        books = books.filter(
            Q(title__icontains=search) |
            Q(authors__name__icontains=search) |
            Q(isbn__icontains=search)
        ).distinct()
    
    if category_id:
        books = books.filter(categories__category_id=category_id)
    
    if status:
        books = books.filter(status=status)
    
    paginator = Paginator(books, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    categories = Category.objects.all()
    
    context = {
        'page_obj': page_obj,
        'categories': categories,
        'search': search,
        'selected_category': category_id,
        'selected_status': status,
    }
    context.update(get_global_stats())
    
    return render(request, 'biblio/book_list.html', context)


# ============================================
# TÉLÉCHARGEMENT ET LECTURE
# ============================================
def download_book(request, book_id):
    """Télécharger le fichier PDF du livre"""
    try:
        book = get_object_or_404(Book, pk=book_id)
        if book.file and book.file.name:
            file_path = book.file.path
            if os.path.exists(file_path):
                file_handle = book.file.open('rb')
                mime_type, encoding = mimetypes.guess_type(file_path)
                if not mime_type:
                    mime_type = 'application/octet-stream'
                
                response = FileResponse(file_handle, content_type=mime_type)
                filename = smart_str(os.path.basename(book.file.name))
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                response['Content-Length'] = book.file.size
                
                return response
        
        raise Http404("Fichier non trouvé")
    except Book.DoesNotExist:
        raise Http404("Livre non trouvé")
    except Exception as e:
        raise Http404(f"Erreur: {str(e)}")


def read_book(request, book_id):
    """Lire le fichier PDF dans le navigateur"""
    try:
        book = get_object_or_404(Book, pk=book_id)
        if book.file and book.file.name:
            file_path = book.file.path
            if os.path.exists(file_path):
                file_handle = book.file.open('rb')
                response = FileResponse(file_handle, content_type='application/pdf')
                filename = smart_str(os.path.basename(book.file.name))
                response['Content-Disposition'] = f'inline; filename="{filename}"'
                response['Content-Length'] = book.file.size
                
                return response
        
        raise Http404("Fichier non trouvé")
    except Book.DoesNotExist:
        raise Http404("Livre non trouvé")
    except Exception as e:
        raise Http404(f"Erreur: {str(e)}")
    



# ============================================
# FONCTION UTILITAIRE POUR OBTENIR LES LIVRES FILTRÉS
# ============================================
def get_filtered_books(request):
    """Récupère les livres avec les mêmes filtres que book_list"""
    books = Book.objects.all().select_related('publisher').prefetch_related('authors', 'categories')
    
    search = request.GET.get('search', '')
    category_id = request.GET.get('category', '')
    status = request.GET.get('status', '')
    
    if search:
        books = books.filter(
            Q(title__icontains=search) |
            Q(authors__name__icontains=search) |
            Q(isbn__icontains=search)
        ).distinct()
    
    if category_id:
        books = books.filter(categories__category_id=category_id)
    
    if status:
        books = books.filter(status=status)
    
    return books


# ============================================
# EXPORT EXCEL
# ============================================
@login_required
def export_books_excel(request):
    """Exporte la liste des livres en Excel"""
    if not openpyxl:
        messages.error(request, 'La bibliothèque openpyxl n\'est pas installée.')
        return redirect('book_list')
    
    # Récupérer les livres filtrés
    books = get_filtered_books(request)
    
    # Créer un classeur Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Liste des Livres"
    
    # Styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # En-têtes
    headers = ['Titre', 'ISBN', 'Auteurs', 'Catégories', 'Éditeur', 'Année', 'Statut', 'Total', 'Disponibles', 'Langue']
    ws.append(headers)
    
    # Appliquer le style aux en-têtes
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Données
    for book in books:
        # Récupérer les auteurs
        authors = ', '.join([f"{a.name}" for a in book.authors.all()])
        if not authors:
            authors = "Aucun auteur"
        
        # Récupérer les catégories
        categories = ', '.join([c.category_name for c in book.categories.all()])
        if not categories:
            categories = "Aucune catégorie"
        
        # Statut en français
        status_map = {
            'available': 'Disponible',
            'borrowed': 'Emprunté',
            'reserved': 'Réservé',
            'maintenance': 'Maintenance'
        }
        status = status_map.get(book.status, book.status)
        
        row = [
            book.title,
            book.isbn or 'N/A',
            authors,
            categories,
            book.publisher.publisher_name if book.publisher else 'N/A',
            book.publication_year or 'N/A',
            status,
            book.total_copies,
            book.available_copies,
            book.language or 'N/A'
        ]
        ws.append(row)
    
    # Ajuster la largeur des colonnes
    for col_num in range(1, len(headers) + 1):
        column_letter = get_column_letter(col_num)
        max_length = 0
        for cell in ws[column_letter]:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Appliquer les bordures à toutes les cellules
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
    
    # Ajouter une ligne de résumé
    ws.append([])
    summary_row = ws.max_row + 1
    ws[f'A{summary_row}'] = f'Total: {books.count()} livres'
    ws[f'A{summary_row}'].font = Font(bold=True, size=11)
    
    # Créer la réponse HTTP
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f'livres_bibliotheque_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Sauvegarder le classeur dans la réponse
    wb.save(response)
    return response


# ============================================
# EXPORT PDF
# ============================================
@login_required
def export_books_pdf(request):
    """Exporte la liste des livres en PDF"""
    if not SimpleDocTemplate:
        messages.error(request, 'La bibliothèque reportlab n\'est pas installée.')
        return redirect('book_list')
    
    # Récupérer les livres filtrés
    books = get_filtered_books(request)
    
    # Créer le PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    # Conteneur pour les éléments
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=1  # Centre
    )
    
    # Titre
    title = Paragraph("Liste des Livres - Bibliothèque", title_style)
    elements.append(title)
    
    # Date
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        alignment=1
    )
    date_text = Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", date_style)
    elements.append(date_text)
    elements.append(Spacer(1, 20))
    
    # Données du tableau
    data = [['Titre', 'ISBN', 'Auteurs', 'Catégories', 'Statut', 'Total', 'Dispo.']]
    
    for book in books:
        # Auteurs
        authors = ', '.join([f"{a.name}" for a in book.authors.all()])[:50]
        if not authors:
            authors = "Aucun"
        
        # Catégories
        categories = ', '.join([c.category_name for c in book.categories.all()])[:40]
        if not categories:
            categories = "Aucune"
        
        # Statut
        status_map = {
            'available': 'Disponible',
            'borrowed': 'Emprunté',
            'reserved': 'Réservé',
            'maintenance': 'Maintenance'
        }
        status = status_map.get(book.status, book.status)
        
        data.append([
            book.title[:40],
            book.isbn or 'N/A',
            authors,
            categories,
            status,
            str(book.total_copies),
            str(book.available_copies)
        ])
    
    # Créer le tableau
    table = Table(data, colWidths=[2.5*inch, 1*inch, 1.5*inch, 1.3*inch, 1*inch, 0.6*inch, 0.6*inch])
    
    # Style du tableau
    table.setStyle(TableStyle([
        # En-tête
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        
        # Corps du tableau
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        
        # Grille
        ('GRID', (0, 0), (-1, -1), 1, colors.white),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#4472C4')),
        
        # Alternance de couleurs
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
        
        # Alignement des nombres
        ('ALIGN', (5, 1), (-1, -1), 'CENTER'),
    ]))
    
    elements.append(table)
    
    # Résumé
    elements.append(Spacer(1, 20))
    summary_style = ParagraphStyle(
        'Summary',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#1e40af'),
    )
    summary = Paragraph(f"<b>Total: {books.count()} livres</b>", summary_style)
    elements.append(summary)
    
    # Construire le PDF
    doc.build(elements)
    
    # Créer la réponse HTTP
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    filename = f'livres_bibliotheque_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


# ============================================
# EXPORT WORD
# ============================================
@login_required
def export_books_word(request):
    """Exporte la liste des livres en Word"""
    if not Document:
        messages.error(request, 'La bibliothèque python-docx n\'est pas installée.')
        return redirect('book_list')
    
    # Récupérer les livres filtrés
    books = get_filtered_books(request)
    
    # Créer un document Word
    doc = Document()
    
    # Titre
    title = doc.add_heading('Liste des Livres - Bibliothèque', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Espace
    
    # Créer le tableau
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # En-têtes
    headers = ['Titre', 'ISBN', 'Auteurs', 'Catégories', 'Statut', 'Total', 'Disponibles']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Style de l'en-tête
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 78, 121)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Données
    for book in books:
        row_cells = table.add_row().cells
        
        # Titre
        row_cells[0].text = book.title
        
        # ISBN
        row_cells[1].text = book.isbn or 'N/A'
        
        # Auteurs
        authors = ', '.join([f"{a.name}" for a in book.authors.all()])
        row_cells[2].text = authors if authors else "Aucun auteur"
        
        # Catégories
        categories = ', '.join([c.category_name for c in book.categories.all()])
        row_cells[3].text = categories if categories else "Aucune catégorie"
        
        # Statut
        status_map = {
            'available': 'Disponible',
            'borrowed': 'Emprunté',
            'reserved': 'Réservé',
            'maintenance': 'Maintenance'
        }
        row_cells[4].text = status_map.get(book.status, book.status)
        
        # Exemplaires
        row_cells[5].text = str(book.total_copies)
        row_cells[6].text = str(book.available_copies)
        
        # Centrer les nombres
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajuster la largeur des colonnes
    widths = [Inches(2.5), Inches(1.2), Inches(1.5), Inches(1.3), Inches(1), Inches(0.7), Inches(0.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    # Résumé
    doc.add_paragraph()
    summary = doc.add_paragraph(f'Total: {books.count()} livres')
    summary.runs[0].font.bold = True
    summary.runs[0].font.size = Pt(12)
    summary.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Créer la réponse HTTP
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f'livres_bibliotheque_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response